"""
Parses a Word (.docx) file into a list of (question, answer) pairs.

Supports two layouts, and will use whichever one it finds content in:

1. TABLE layout (recommended - most reliable):
   A table with two columns. First row can optionally be a header
   like "Question" / "Answer" - it's auto-detected and skipped.

   | Question                | Answer                                    |
   |--------------------------|-------------------------------------------|
   | What is a BVS device?    | BVS stands for ...                        |
   | What does ARM mean?      | ARM refers to ...                         |

2. Q:/A: PARAGRAPH layout:
   Plain paragraphs, alternating, each starting with "Q:" and "A:"
   (case-insensitive, colon required):

   Q: What is a BVS device?
   A: BVS stands for ...

   Q: What does ARM mean?
   A: ARM refers to ...

If a document has both, table pairs are loaded first, then paragraph
pairs are appended.
"""
from docx import Document


def _clean(text: str) -> str:
    return " ".join(text.split()).strip()


def _looks_like_header(question_cell: str, answer_cell: str) -> bool:
    q = question_cell.strip().lower()
    a = answer_cell.strip().lower()
    return q in {"question", "questions", "q"} and a in {"answer", "answers", "a"}


def _load_from_tables(doc: Document):
    pairs = []
    for table in doc.tables:
        for i, row in enumerate(table.rows):
            cells = [c.text for c in row.cells]
            if len(cells) < 2:
                continue
            question, answer = _clean(cells[0]), _clean(cells[1])
            if not question or not answer:
                continue
            if i == 0 and _looks_like_header(question, answer):
                continue
            pairs.append((question, answer))
    return pairs


def _load_from_paragraphs(doc: Document):
    pairs = []
    pending_question = None
    for para in doc.paragraphs:
        text = _clean(para.text)
        if not text:
            continue
        lower = text.lower()
        if lower.startswith("q:") or lower.startswith("q :"):
            pending_question = text.split(":", 1)[1].strip()
        elif (lower.startswith("a:") or lower.startswith("a :")) and pending_question:
            answer = text.split(":", 1)[1].strip()
            pairs.append((pending_question, answer))
            pending_question = None
    return pairs


def parse_qa_docx(path: str):
    """
    Returns a list of (question, answer) tuples found in the document.
    Raises FileNotFoundError / docx errors upward - caller should handle.
    """
    doc = Document(path)

    pairs = _load_from_tables(doc)
    pairs.extend(_load_from_paragraphs(doc))

    if not pairs:
        raise ValueError(
            "No Q&A pairs found. Use either a two-column table (Question | Answer) "
            "or paragraphs starting with 'Q:' and 'A:'."
        )
    return pairs
